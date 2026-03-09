"""
Markdown to Word (.docx) converter for project documentation.
Converts all key markdown files in Phase 1 to formatted Word documents.
"""
import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

SRC_DIR = r"E:\meiz\zhuangce-ai\zhuangce-ai-最终归档-第一批到第八批\第一批-项目立项层-阶段1"
OUT_DIR = os.path.join(SRC_DIR, "Word版")
os.makedirs(OUT_DIR, exist_ok=True)

FILES = [
    "A-第一阶段成果总览.md",
    "B-数据分析与算法成果报告.md",
    "C-可视化图表索引.md",
    "D-第二阶段推进计划.md",
    "01-校赛重点品类与产品确认.md",
    "02-产品预测目标文档-示例正式版.md",
    "22-第一批实操开始清单-今天.md",
    "23-第一批实操开始清单-明天.md",
    "24-第一批你先填写的内容模板.md",
]


def parse_table(lines):
    """Parse markdown table lines into rows of cells."""
    rows = []
    for line in lines:
        line = line.strip()
        if line.startswith("|"):
            cells = [c.strip() for c in line.split("|")[1:-1]]
            rows.append(cells)
    # Remove separator row (contains only dashes/colons)
    if len(rows) > 1 and all(re.match(r'^[-:]+$', c) for c in rows[1]):
        rows.pop(1)
    return rows


def add_formatted_run(paragraph, text):
    """Add text with inline bold/code/red formatting to a paragraph."""
    # Split by bold markers, code markers, and [[红:...]] red markers
    parts = re.split(r'(\*\*.*?\*\*|`[^`]+`|\[\[红:.*?\]\])', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)
        elif part.startswith('[[红:') and part.endswith(']]'):
            inner = part[4:-2]
            run = paragraph.add_run(inner)
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        else:
            paragraph.add_run(part)


def md_to_docx(md_path, docx_path):
    """Convert a markdown file to a Word document with formatting."""
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    
    # Configure heading styles
    for level in range(1, 5):
        hstyle = doc.styles[f'Heading {level}']
        hstyle.font.name = '微软雅黑'
        hstyle.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    i = 0
    in_code_block = False
    code_lines = []
    
    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip('\n')
        
        # Code block handling
        if stripped.strip().startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                pf = p.paragraph_format
                pf.left_indent = Inches(0.3)
                pf.space_before = Pt(4)
                pf.space_after = Pt(4)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(stripped)
            i += 1
            continue
        
        # Headings
        heading_match = re.match(r'^(#{1,4})\s+(.*)', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            h = doc.add_heading(level=level)
            add_formatted_run(h, text)
            i += 1
            continue
        
        # Table detection
        if stripped.strip().startswith('|') and i + 1 < len(lines) and lines[i + 1].strip().startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            rows = parse_table(table_lines)
            if rows:
                num_cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=num_cols, style='Light Grid Accent 1')
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for ri, row in enumerate(rows):
                    for ci, cell_text in enumerate(row):
                        if ci < num_cols:
                            cell = table.cell(ri, ci)
                            cell.text = ''
                            p = cell.paragraphs[0]
                            add_formatted_run(p, cell_text)
                            if ri == 0:
                                for run in p.runs:
                                    run.bold = True
            continue
        
        # Horizontal rule
        if re.match(r'^---+\s*$', stripped):
            # Add a thin horizontal line via paragraph border
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue
        
        # Blockquote
        if stripped.strip().startswith('>'):
            text = re.sub(r'^>\s*', '', stripped.strip())
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            run = p.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue
        
        # Unordered list
        ul_match = re.match(r'^(\s*)[-*]\s+(.*)', stripped)
        if ul_match:
            indent = len(ul_match.group(1))
            text = ul_match.group(2)
            p = doc.add_paragraph(style='List Bullet')
            if indent >= 2:
                p.paragraph_format.left_indent = Inches(0.3 * (indent // 2 + 1))
            add_formatted_run(p, text)
            i += 1
            continue
        
        # Ordered list
        ol_match = re.match(r'^(\s*)\d+\.\s+(.*)', stripped)
        if ol_match:
            indent = len(ol_match.group(1))
            text = ol_match.group(2)
            p = doc.add_paragraph(style='List Number')
            if indent >= 2:
                p.paragraph_format.left_indent = Inches(0.3 * (indent // 2 + 1))
            add_formatted_run(p, text)
            i += 1
            continue
        
        # Empty line
        if stripped.strip() == '':
            i += 1
            continue
        
        # Normal paragraph
        p = doc.add_paragraph()
        add_formatted_run(p, stripped.strip())
        i += 1
    
    doc.save(docx_path)
    return docx_path


# Convert all files
print(f"Converting {len(FILES)} markdown files to Word...")
for fname in FILES:
    src = os.path.join(SRC_DIR, fname)
    if not os.path.exists(src):
        print(f"  SKIP (not found): {fname}")
        continue
    dst = os.path.join(OUT_DIR, fname.replace('.md', '.docx'))
    try:
        md_to_docx(src, dst)
        print(f"  OK: {fname} -> {os.path.basename(dst)}")
    except Exception as e:
        print(f"  ERROR: {fname} -> {e}")

print(f"\nDone! Word files saved to: {OUT_DIR}")
